import os
import re
import uuid
import json
import sqlite3
import datetime
from flask import Flask, request, render_template, redirect, url_for, flash
from werkzeug.utils import secure_filename
import pdfplumber
import docx
import requests

try:
    import spacy
except ImportError:
    spacy = None

try:
    import openai
except ImportError:
    openai = None

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOAD_DIR = os.path.join(BASE_DIR, "uploads")
DB_PATH = os.path.join(BASE_DIR, "profiles.db")
ALLOWED_EXTENSIONS = {"pdf", "docx"}

app = Flask(__name__)
app.secret_key = "change-this-secret"
app.config["UPLOAD_FOLDER"] = UPLOAD_DIR
os.makedirs(UPLOAD_DIR, exist_ok=True)

SKILL_KEYWORDS = [
    "python", "django", "flask", "javascript", "react", "node", "express", "sql",
    "postgresql", "mysql", "mongodb", "aws", "azure", "gcp", "machine learning",
    "deep learning", "nlp", "data science", "pytorch", "tensorflow", "scikit-learn",
    "html", "css", "git", "docker", "kubernetes", "rest api", "api", "linux",
    "java", "c++", "c#", "ruby", "php", "android", "ios", "swift", "kotlin",
    "pandas", "numpy", "tableau", "power bi", "spark", "hadoop", "unity", "unreal",
    "selenium", "pytest", "unittest", "jira", "ansible", "terraform", "shell"
]

ROLE_RULES = [
    ("Machine Learning Engineer", {"machine learning", "deep learning", "nlp", "tensorflow", "pytorch", "pandas", "numpy"}),
    ("Data Scientist", {"data science", "python", "pandas", "scikit-learn", "machine learning", "numpy"}),
    ("Backend Developer", {"django", "flask", "node", "express", "rest api", "sql", "mongodb", "postgresql", "mysql"}),
    ("Frontend Developer", {"react", "javascript", "html", "css", "vue", "angular", "typescript"}),
    ("Full Stack Developer", {"django", "flask", "react", "javascript", "node", "html", "css"}),
    ("DevOps Engineer", {"docker", "kubernetes", "aws", "azure", "gcp", "linux", "terraform", "ansible"}),
    ("Data Engineer", {"spark", "hadoop", "sql", "python", "aws", "gcp", "etl"}),
    ("QA Engineer", {"automation", "selenium", "pytest", "unittest", "jira"}),
    ("Mobile Developer", {"android", "ios", "swift", "kotlin"})
]

JOB_SEARCH_SITES = {
    "LinkedIn": "https://www.linkedin.com/jobs/search/?keywords={query}",
    "Indeed": "https://www.indeed.com/jobs?q={query}",
    "Glassdoor": "https://www.glassdoor.com/Job/jobs.htm?sc.keyword={query}",
}

SECTION_PATTERNS = {
    "education": re.compile(r"(education|academic background|qualifications|degree|university|college)", re.I),
    "experience": re.compile(r"(experience|work history|professional experience|internship|employment)", re.I),
    "projects": re.compile(r"(projects|academic projects|project experience|portfolio)", re.I),
    "skills": re.compile(r"(skills|technical skills|programming skills|tools)", re.I),
}

BULLET_PATTERN = re.compile(r"^[\-\u2022\*\d\.\)]+\s*(.*)$")

SPACY_MODEL_NAME = "en_core_web_sm"
nlp = None
if spacy:
    try:
        nlp = spacy.load(SPACY_MODEL_NAME)
    except OSError:
        nlp = None

OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY")
OPENAI_AVAILABLE = False
if openai and OPENAI_API_KEY:
    openai.api_key = OPENAI_API_KEY
    OPENAI_AVAILABLE = True


def get_db_connection():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn


def init_db():
    conn = get_db_connection()
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS profiles (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT,
            skills TEXT,
            education TEXT,
            experience TEXT,
            projects TEXT,
            roles TEXT,
            resume_file TEXT,
            cover_letter TEXT,
            created_at TEXT
        )
        """
    )
    conn.commit()
    columns = [row[1] for row in conn.execute("PRAGMA table_info(profiles)").fetchall()]
    if "cover_letter" not in columns:
        conn.execute("ALTER TABLE profiles ADD COLUMN cover_letter TEXT")
        conn.commit()
    conn.close()


def serialize_list(items):
    return json.dumps(items, ensure_ascii=False)


def deserialize_list(value):
    try:
        return json.loads(value) if value else []
    except json.JSONDecodeError:
        return []


def save_profile(name, skills, education, experience, projects, roles, resume_file, cover_letter=None):
    conn = get_db_connection()
    now = datetime.datetime.utcnow().isoformat()
    cursor = conn.execute(
        "INSERT INTO profiles (name, skills, education, experience, projects, roles, resume_file, cover_letter, created_at) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)",
        (
            name,
            serialize_list(skills),
            serialize_list(education),
            serialize_list(experience),
            serialize_list(projects),
            serialize_list(roles),
            resume_file,
            cover_letter,
            now,
        ),
    )
    conn.commit()
    profile_id = cursor.lastrowid
    conn.close()
    return profile_id


def update_profile(profile_id, name, skills, education, experience, projects, roles, cover_letter=None):
    conn = get_db_connection()
    conn.execute(
        "UPDATE profiles SET name = ?, skills = ?, education = ?, experience = ?, projects = ?, roles = ?, cover_letter = ? WHERE id = ?",
        (
            name,
            serialize_list(skills),
            serialize_list(education),
            serialize_list(experience),
            serialize_list(projects),
            serialize_list(roles),
            cover_letter,
            profile_id,
        ),
    )
    conn.commit()
    conn.close()


def list_profiles():
    conn = get_db_connection()
    rows = conn.execute("SELECT id, name, created_at FROM profiles ORDER BY created_at DESC").fetchall()
    conn.close()
    return [dict(row) for row in rows]


def get_profile(profile_id):
    conn = get_db_connection()
    row = conn.execute("SELECT * FROM profiles WHERE id = ?", (profile_id,)).fetchone()
    conn.close()
    if not row:
        return None
    return {
        "id": row["id"],
        "name": row["name"],
        "skills": deserialize_list(row["skills"]),
        "education": deserialize_list(row["education"]),
        "experience": deserialize_list(row["experience"]),
        "projects": deserialize_list(row["projects"]),
        "roles": deserialize_list(row["roles"]),
        "resume_file": row["resume_file"],
        "cover_letter": row["cover_letter"],
        "created_at": row["created_at"],
    }


def init_db():
    conn = get_db_connection()
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS profiles (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT,
            skills TEXT,
            education TEXT,
            experience TEXT,
            projects TEXT,
            roles TEXT,
            resume_file TEXT,
            cover_letter TEXT,
            created_at TEXT
        )
        """
    )
    conn.commit()
    columns = [row[1] for row in conn.execute("PRAGMA table_info(profiles)").fetchall()]
    if "cover_letter" not in columns:
        conn.execute("ALTER TABLE profiles ADD COLUMN cover_letter TEXT")
        conn.commit()
    conn.close()


init_db()


def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def extract_text_from_pdf(path):
    text = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text.append(page_text)
    return "\n".join(text)


def extract_text_from_docx(path):
    document = docx.Document(path)
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return "\n".join(paragraphs)


def parse_resume(path):
    ext = path.rsplit('.', 1)[1].lower()
    if ext == 'pdf':
        return extract_text_from_pdf(path)
    if ext == 'docx':
        return extract_text_from_docx(path)
    return ""


def extract_skills(text):
    normalized = text.lower()
    found = set()
    for keyword in SKILL_KEYWORDS:
        pattern = r"\b" + re.escape(keyword) + r"\b"
        if re.search(pattern, normalized):
            found.add(keyword.title())
    return sorted(found)


def extract_sentences(text):
    if nlp:
        doc = nlp(text)
        return [sent.text.strip() for sent in doc.sents if sent.text.strip()]
    return [line.strip() for line in text.splitlines() if line.strip()]


def extract_text_sections(text):
    sentences = extract_sentences(text)
    sections = {"education": [], "experience": [], "projects": [], "skills": []}
    current = None
    for sentence in sentences:
        header = sentence.lower().strip(':')
        section_found = False
        for section, pattern in SECTION_PATTERNS.items():
            if pattern.search(header) and len(sentence) < 120:
                current = section
                section_found = True
                break
        if section_found:
            continue
        clean_text = sentence
        bullet_match = BULLET_PATTERN.match(sentence)
        if bullet_match:
            clean_text = bullet_match.group(1).strip()
        if current:
            sections[current].append(clean_text)
        else:
            lowered = clean_text.lower()
            if any(keyword in lowered for keyword in ["university", "college", "degree", "bachelor", "master", "phd"]):
                sections["education"].append(clean_text)
            elif any(keyword in lowered for keyword in ["intern", "engineer", "developer", "analyst", "manager", "consultant"]):
                sections["experience"].append(clean_text)
            elif any(keyword in lowered for keyword in ["project", "built", "designed", "developed", "created"]):
                sections["projects"].append(clean_text)
            elif any(skill in lowered for skill in SKILL_KEYWORDS):
                sections["skills"].append(clean_text)
    return {
        "education": sections["education"][:10],
        "experience": sections["experience"][:10],
        "projects": sections["projects"][:10],
        "skills": sections["skills"][:12],
    }


def extract_name(text):
    if nlp:
        doc = nlp(text)
        for ent in doc.ents:
            if ent.label_ == "PERSON" and len(ent.text.split()) <= 3:
                return ent.text
    return None


def generate_cover_letter(profile_name, top_role, skills, projects, experience, company_name=None):
    experience_summary = " ".join(experience[:2])
    project_summary = projects[0] if projects else "relevant projects"
    if OPENAI_AVAILABLE:
        prompt = (
            f"Write a professional and concise cover letter for a {top_role} role. "
            f"Use the following resume details:\n"
            f"Name: {profile_name}\n"
            f"Skills: {', '.join(skills)}\n"
            f"Key experience: {experience_summary}\n"
            f"Recent project: {project_summary}\n"
            f"Company: {company_name or 'a forward-looking employer'}\n"
            f"Keep it a friendly, confident, single-page letter."
        )
        try:
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "You are a skilled cover letter writer."},
                    {"role": "user", "content": prompt},
                ],
                max_tokens=350,
                temperature=0.7,
            )
            return response.choices[0].message.content.strip()
        except Exception:
            pass

    skill_list = ", ".join(skills[:6]) if skills else "relevant technical and problem-solving skills"
    return (
        f"Dear Hiring Manager,\n\n"
        f"I am excited to apply for the {top_role} position. With experience working on {project_summary} and proficiency in {skill_list}, "
        f"I can contribute to your team by delivering strong technical solutions and a thoughtful approach to product development. "
        f"My background includes {experience_summary or 'relevant internship and project experience'}, "
        f"and I enjoy collaborating to solve complex problems.\n\n"
        f"Thank you for considering my application. I look forward to the opportunity to discuss how my experience and skills align with your needs.\n\n"
        f"Sincerely,\n{profile_name if profile_name else 'A motivated applicant'}"
    )


def compute_role_scores(skills):
    normalized = {s.lower() for s in skills}
    scored = []
    for role, keywords in ROLE_RULES:
        matched = normalized & keywords
        if not matched:
            continue
        score = round(100 * len(matched) / len(keywords))
        scored.append({
            "role": role,
            "score": score,
            "matches": sorted(matched),
            "description": f"Matched {len(matched)} of {len(keywords)} keywords"
        })
    scored.sort(key=lambda x: (-x["score"], x["role"]))
    return scored


def recommend_roles(skills):
    scores = compute_role_scores(skills)
    if not scores:
        return [{"role": "General Software Engineer", "score": 50, "matches": [], "description": "No strong role match detected"}]
    return scores[:5]

MARKET_SKILL_DEMAND = {
    "python": 95,
    "docker": 90,
    "rest api": 88,
    "sql": 85,
    "javascript": 93,
    "react": 87,
    "aws": 92,
    "kubernetes": 84,
    "machine learning": 89,
    "nlp": 78,
    "tensorflow": 80,
    "pytorch": 79,
    "postgresql": 81,
    "mongodb": 76,
    "git": 94,
    "linux": 83,
    "terraform": 74,
    "ansible": 70,
    "spark": 75,
    "hadoop": 65,
}

COMPANY_HIRING_MAP = {
    "Backend Developer": ["Amazon", "Microsoft", "Stripe", "Square", "Salesforce"],
    "Frontend Developer": ["Meta", "Shopify", "Adobe", "Spotify", "Airbnb"],
    "Full Stack Developer": ["HubSpot", "Coinbase", "Twilio", "Atlassian", "Zendesk"],
    "DevOps Engineer": ["Google", "Netflix", "Datadog", "HashiCorp", "Red Hat"],
    "Machine Learning Engineer": ["OpenAI", "NVIDIA", "Apple", "Google", "IBM"],
    "Data Scientist": ["Uber", "Airbnb", "Capital One", "LinkedIn", "Netflix"],
    "Data Engineer": ["Snowflake", "Databricks", "Oracle", "Palantir", "Cloudera"],
    "QA Engineer": ["Intel", "Cisco", "Microsoft", "Qualcomm", "Palo Alto Networks"],
    "Mobile Developer": ["Google", "Apple", "Uber", "Snapchat", "Spotify"],
}

ROLE_BASE_SALARY = {
    "Backend Developer": 85000,
    "Frontend Developer": 78000,
    "Full Stack Developer": 82000,
    "DevOps Engineer": 90000,
    "Machine Learning Engineer": 95000,
    "Data Scientist": 92000,
    "Data Engineer": 88000,
    "QA Engineer": 70000,
    "Mobile Developer": 78000,
    "General Software Engineer": 76000,
}


def build_career_insights(skills, role_scores):
    top_role = role_scores[0]["role"] if role_scores else "Software Engineer"
    top_score = role_scores[0]["score"] if role_scores else 50
    normalized = {s.lower() for s in skills}
    role_keywords = next((keywords for role, keywords in ROLE_RULES if role == top_role), set())
    missing = sorted(role_keywords - normalized)
    missing_skills = [keyword.title() for keyword in missing[:3]]
    readiness = top_score
    improved = min(100, readiness + 20) if missing_skills else readiness

    companies = COMPANY_HIRING_MAP.get(top_role, ["Growing startups", "regional tech firms", "remote-first teams"])
    expected_salary = ROLE_BASE_SALARY.get(top_role, ROLE_BASE_SALARY["General Software Engineer"])
    salary_growth = {
        "3_months": int(expected_salary * (1 + max(0, readiness - 50) / 1200 + 0.02)),
        "1_year": int(expected_salary * (1 + max(0, readiness - 50) / 900 + 0.08)),
        "3_years": int(expected_salary * (1 + max(0, readiness - 50) / 700 + 0.20)),
    }

    skill_demand = []
    for skill in sorted(set(skills), key=lambda x: (-MARKET_SKILL_DEMAND.get(x.lower(), 50), x))[:10]:
        skill_demand.append({
            "skill": skill,
            "market_demand": MARKET_SKILL_DEMAND.get(skill.lower(), 60),
            "your_level": min(95, 50 + int(len(skill) * 2)) if skill else 50,
        })

    return {
        "top_role": top_role,
        "readiness": readiness,
        "missing_skills": missing_skills,
        "improved_readiness": improved,
        "salary_growth": salary_growth,
        "companies": companies[:5],
        "skill_demand": skill_demand,
    }


def build_recruiter_review(text, skills, experience, projects):
    normalized = text.lower()
    num_skills = len(skills)
    leadership_keywords = ["lead", "managed", "mentor", "owned", "collaborated", "team", "strategy"]
    leader_count = sum(normalized.count(word) for word in leadership_keywords)
    tech_keywords = ["developed", "designed", "built", "implemented", "deployed", "optimized", "scaled"]
    tech_count = sum(normalized.count(word) for word in tech_keywords)
    bullet_count = len(re.findall(r"^[\-\u2022\*\d\.\)]+", text, re.M))

    confidence = min(100, 35 + num_skills * 3 + bullet_count * 2)
    leadership = min(100, 20 + leader_count * 8 + (len(projects) * 5))
    technical_depth = min(100, 30 + tech_count * 4 + num_skills * 2 + len(experience) * 3)
    communication = min(100, 40 + min(20, max(0, 90 - len(text) // 200)))
    ats = min(100, 35 + num_skills * 3 + bullet_count * 2 + int(bool(skills)) * 10)
    portfolio = min(100, 30 + len(projects) * 10 + int("project" in normalized) * 5)

    low_scores = sorted([
        (confidence, "confidence"),
        (leadership, "leadership"),
        (technical_depth, "technical depth"),
        (communication, "communication"),
        (ats, "ATS compatibility"),
        (portfolio, "portfolio quality"),
    ])
    lowest_score, lowest_area = low_scores[0]

    if lowest_area == "technical depth":
        reason = "The resume could use stronger technical depth: include specific frameworks, architecture decisions, metrics, and system ownership."
        tip = "Add concrete achievements like 'reduced latency by 30%' or 'deployed a microservice using Docker + AWS'."
    elif lowest_area == "leadership":
        reason = "Leadership and collaboration examples are light. Recruiters look for ownership, mentoring, and cross-team communication."
        tip = "Mention team projects, leadership roles, or examples where you guided peers and delivered results."
    elif lowest_area == "communication":
        reason = "The resume may feel dense or hard to scan. Clear formatting and concise bullets improve recruiter confidence."
        tip = "Use short, action-oriented bullet points and avoid long paragraphs in experience descriptions."
    elif lowest_area == "ats compatibility":
        reason = "The resume may not be optimized for applicant tracking systems. Use plain section headings and keyword-rich phrases."
        tip = "Keep formatting simple, use standard headings like 'Experience' and 'Skills', and avoid tables or images in the resume."
    else:
        reason = "This resume is on the right track, but it can still stand out by adding one or two higher-impact examples."
        tip = "Highlight top achievements, quantify outcomes, and align your resume with the role you want."

    return {
        "confidence": confidence,
        "leadership": leadership,
        "technical_depth": technical_depth,
        "communication": communication,
        "ats": ats,
        "portfolio": portfolio,
        "reason": reason,
        "tip": tip,
    }


def build_summarized_feedback(career_insights, recruiter_review):
    if career_insights["missing_skills"]:
        missing_text = ", ".join(career_insights["missing_skills"])
        readiness_text = (
            f"You are currently {career_insights['readiness']}% ready for {career_insights['top_role']} roles. "
            f"Learn {missing_text} → readiness becomes {career_insights['improved_readiness']}%."
        )
    else:
        readiness_text = (
            f"You are currently {career_insights['readiness']}% ready for {career_insights['top_role']} roles. "
            f"Your current skill mix already aligns well with core requirements."
        )

    return {
        "overview": readiness_text,
        "rejection": recruiter_review["reason"],
        "quick_tip": recruiter_review["tip"],
    }


def build_profile_insights(profile):
    text = "\n".join(profile.get(field, []) for field in ["education", "experience", "projects"]) if isinstance(profile.get("education"), list) else ""
    career_insights = build_career_insights(profile.get("skills", []), recommend_roles(profile.get("skills", [])))
    recruiter_review = build_recruiter_review(text, profile.get("skills", []), profile.get("experience", []), profile.get("projects", []))
    feedback = build_summarized_feedback(career_insights, recruiter_review)
    return career_insights, recruiter_review, feedback


def build_job_search_links(roles, skills):
    queries = []
    if roles:
        queries.append(" ".join(role["role"] for role in roles[:2]))
    if skills:
        queries.append(" ".join(skills[:4]))
    queries.append("software engineer internship")
    queries = [q.strip().replace(' ', '+') for q in queries if q.strip()]
    search_links = []
    for query in queries:
        for site, url in JOB_SEARCH_SITES.items():
            search_links.append({
                "site": site,
                "query": query.replace('+', ' '),
                "url": url.format(query=query)
            })
    return search_links


def search_jobs_api(query, limit=8):
    if not query:
        return []
    query = query.strip()
    jobs = []
    try:
        response = requests.get("https://remotive.io/api/remote-jobs", params={"search": query}, timeout=8)
        response.raise_for_status()
        data = response.json()
        for job in data.get("jobs", [])[:limit]:
            jobs.append({
                "title": job.get("title", "Unknown title"),
                "company": job.get("company_name", "Unknown company"),
                "location": job.get("candidate_required_location", "Remote"),
                "url": job.get("url"),
                "source": "Remotive",
            })
    except Exception:
        jobs = []
    if not jobs:
        try:
            response = requests.get("https://www.arbeitnow.com/api/job-board-api", timeout=8)
            response.raise_for_status()
            data = response.json()
            for item in data.get("data", []):
                title = item.get("title", "")
                company = item.get("company_name", "")
                description = item.get("description", "").lower()
                if query.lower() in title.lower() or query.lower() in description:
                    jobs.append({
                        "title": title,
                        "company": company or item.get("location", ""),
                        "location": item.get("location", "Remote"),
                        "url": item.get("url"),
                        "source": "ArbeitNow",
                    })
                    if len(jobs) >= limit:
                        break
        except Exception:
            return []
    return jobs[:limit]


@app.route('/', methods=['GET', 'POST'])
def index():
    return render_template('index.html')


@app.route('/analyze', methods=['POST'])
def analyze():
    if 'resume' not in request.files:
        flash('No resume file uploaded.', 'error')
        return redirect(url_for('index'))
    file = request.files['resume']
    if file.filename == '':
        flash('Please select a file.', 'error')
        return redirect(url_for('index'))
    if not allowed_file(file.filename):
        flash('Allowed file types are PDF and DOCX only.', 'error')
        return redirect(url_for('index'))

    filename = secure_filename(file.filename)
    unique_name = f"{uuid.uuid4().hex}_{filename}"
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], unique_name)
    file.save(filepath)

    text = parse_resume(filepath)
    skills = extract_skills(text)
    sections = extract_text_sections(text)
    role_scores = recommend_roles(skills)
    roles = [item["role"] for item in role_scores]
    job_links = build_job_search_links(role_scores, skills)
    api_jobs = search_jobs_api(roles[0] if roles else "software engineer")
    cover_letter = generate_cover_letter(
        profile_name=request.form.get('profile_name', '').strip() or 'Applicant',
        top_role=roles[0] if roles else 'Software Engineer',
        skills=skills,
        projects=sections['projects'],
        experience=sections['experience'],
    )

    career_insights = build_career_insights(skills, role_scores)
    recruiter_review = build_recruiter_review(text, skills, sections['experience'], sections['projects'])
    feedback = build_summarized_feedback(career_insights, recruiter_review)

    profile_name = request.form.get('profile_name', '').strip() or 'Unnamed Profile'
    save_profile_flag = request.form.get('save_profile') == 'on'
    profile_id = None
    if save_profile_flag:
        profile_id = save_profile(
            profile_name,
            skills,
            sections['education'],
            sections['experience'],
            sections['projects'],
            roles,
            unique_name,
            cover_letter=cover_letter,
        )

    return render_template(
        'results.html',
        skills=skills,
        education=sections['education'],
        experience=sections['experience'],
        projects=sections['projects'],
        roles=roles,
        role_scores=role_scores,
        job_links=job_links,
        api_jobs=api_jobs,
        cover_letter=cover_letter,
        raw_text=text[:1500] + '...' if len(text) > 1500 else text,
        profile_saved=bool(profile_id),
        profile_id=profile_id,
        profile_name=profile_name,
        career_insights=career_insights,
        recruiter_review=recruiter_review,
        feedback=feedback,
    )


@app.route('/profiles')
def profiles():
    profiles = list_profiles()
    return render_template('profiles.html', profiles=profiles)


@app.route('/profile/<int:profile_id>')
def profile_detail(profile_id):
    profile = get_profile(profile_id)
    if not profile:
        flash('Profile not found.', 'error')
        return redirect(url_for('profiles'))
    career_insights, recruiter_review, feedback = build_profile_insights(profile)
    return render_template(
        'profile.html',
        profile=profile,
        career_insights=career_insights,
        recruiter_review=recruiter_review,
        feedback=feedback,
    )


@app.route('/profile/<int:profile_id>/edit', methods=['GET', 'POST'])
def edit_profile(profile_id):
    profile = get_profile(profile_id)
    if not profile:
        flash('Profile not found.', 'error')
        return redirect(url_for('profiles'))

    if request.method == 'POST':
        name = request.form.get('name', profile['name']).strip() or profile['name']
        skills = [line.strip() for line in request.form.get('skills', '').splitlines() if line.strip()]
        education = [line.strip() for line in request.form.get('education', '').splitlines() if line.strip()]
        experience = [line.strip() for line in request.form.get('experience', '').splitlines() if line.strip()]
        projects = [line.strip() for line in request.form.get('projects', '').splitlines() if line.strip()]
        roles = [line.strip() for line in request.form.get('roles', '').splitlines() if line.strip()]
        cover_letter = request.form.get('cover_letter', '').strip() or None

        if not roles:
            roles = [item['role'] for item in recommend_roles(skills)]

        update_profile(profile_id, name, skills, education, experience, projects, roles, cover_letter=cover_letter)
        flash('Profile updated successfully.', 'success')
        return redirect(url_for('profile_detail', profile_id=profile_id))

    return render_template('edit_profile.html', profile=profile)


if __name__ == '__main__':
    app.run(debug=True)
